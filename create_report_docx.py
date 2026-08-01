import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def add_callout(doc, text, title="TECHNICAL NOTE", border_color="00F2FE", fill_color="F0FDFF"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, fill_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:top w:val="none"/>
            <w:right w:val="none"/>
            <w:bottom w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"📌 {title}: ")
    r_title.bold = True
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = RGBColor(11, 19, 43)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Arial'
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def generate_report():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles & Fonts
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)
    
    # Colors
    NAVY = RGBColor(11, 19, 43)
    CYAN = RGBColor(0, 136, 169)
    DARK_BLUE = RGBColor(28, 37, 65)
    GRAY = RGBColor(100, 116, 139)
    
    # =========================================================================
    # COVER / TITLE SECTION
    # =========================================================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("RIDE PULSE")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(32)
    run_title.bold = True
    run_title.font.color.rgb = CYAN

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)
    run_sub = p_sub.add_run("Multi-Feature Bike Tracking Application with SOS & Off-Grid Mesh Network")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(16)
    run_sub.bold = True
    run_sub.font.color.rgb = NAVY

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_after = Pt(30)
    run_desc = p_desc.add_run("Comprehensive Software Architecture, Technical Design Specification, and Function-by-Function In-Depth Code Reference Report")
    run_desc.font.name = 'Arial'
    run_desc.font.size = Pt(11)
    run_desc.italic = True
    run_desc.font.color.rgb = GRAY

    # Metadata Box Table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_table, "00F2FE")
    meta_data = [
        ("Course / Degree:", "MSc IT (Part-3) Final Technical Project Report"),
        ("Project Name:", "RidePulse (Bike Telemetry, Digital Garage & SOS Mesh Network)"),
        ("Technology Stack:", "HTML5, CSS3 Glassmorphism, JavaScript (ES6+), Leaflet, Web Audio API, Web Speech API, IndexedDB, Canvas 2D, OSRM"),
        ("Document Version:", "1.0.0 (Production Release)"),
        ("Report Scope:", "Full Codebase Function Inventory, Data Models & Architecture")
    ]
    for idx, (label, val) in enumerate(meta_data):
        c0 = meta_table.cell(idx, 0)
        c1 = meta_table.cell(idx, 1)
        set_cell_background(c0, "F8FAFC")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, 80, 80, 120, 120)
        set_cell_margins(c1, 80, 80, 120, 120)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = NAVY
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = DARK_BLUE
        
    doc.add_page_break()

    # Helper function for Headings
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(18)
        r.bold = True
        r.font.color.rgb = NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = CYAN
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(11.5)
        r.bold = True
        r.font.color.rgb = DARK_BLUE
        return p

    def add_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.color.rgb = NAVY
        r = p.add_run(text)
        return p

    # =========================================================================
    # 1. EXECUTIVE SUMMARY & SYSTEM OVERVIEW
    # =========================================================================
    add_h1("1. Executive Summary & System Overview")
    
    add_p("RidePulse is an advanced, multi-feature web application engineered specifically for motorcycle riders, touring crews, and adventure cyclists. Standard navigation and fitness applications fail to address critical motorcycle-specific challenges—such as riding through remote cellular dead zones, managing multi-vehicle maintenance schedules, coordinating group rides, detecting high-G crash impacts, and delivering hands-free emergency SOS alerts.")
    
    add_p("RidePulse resolves these vulnerabilities by combining modular JavaScript software engines with browser-native APIs (Web Audio API, Web Speech API, Geolocation, IndexedDB, and HTML5 Canvas 2D). It operates completely client-side without external backend servers, ensuring zero latency, instant response, and off-grid reliability.")

    add_callout(doc, "RidePulse incorporates 10 core subsystems: Digital Garage, Group Ride Hub & SVG QR Invites, Telemetry HUD & Lean Angle Engine, Off-Grid P2P Mesh Network Simulator, Voice Intercom, Automated Guardian Safety Check-In, Geotagged Community Hazard Engine, Dual-Mode SOS Emergency Dispatch, Multi-Layer Google Navigation, and Ride Analytics with GPX Track Exporting.", "SYSTEM HIGHLIGHT")

    add_h2("1.1 Problem Statement vs. RidePulse Solutions")
    
    prob_table = doc.add_table(rows=6, cols=3)
    prob_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(prob_table, "CBD5E1")
    
    headers = ["Industry Challenge / Problem", "Risk to Riders", "RidePulse Architectural Solution"]
    for c_idx, h_text in enumerate(headers):
        cell = prob_table.cell(0, c_idx)
        set_cell_background(cell, "0B132B")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    problems_data = [
        ("Cellular Signal Loss in Remote Twisty Roads", "Riders lose live location sharing and cannot call for emergency help in dead zones.", "Off-Grid P2P Mesh Network Protocol (`mesh.js`) simulating hop-by-hop packet relay across crew devices."),
        ("Unnoticed Mechanical Wear & Overdue Maintenance", "Chain snap, low tire tread, or engine seizure during high-speed highway rides.", "Digital Garage Engine (`garage.js`) tracking mileage-based service intervals and pre-ride safety gates."),
        ("Rider Unresponsiveness / Solo Crash Isolation", "A solo rider crashes in a ravine and remains unconscious without anyone knowing.", "Automated Guardian Safety Check-In (`guardian.js`) with timed prompts and automatic SOS timeout dispatch."),
        ("Blind Corner Hazards (Potholes, Oil Spills)", "Subsequent group riders hit unexpected road hazards at high cornering speeds.", "Community Hazard Reporting Engine (`hazards.js`) with IndexedDB storage, map pins, and mesh broadcasts."),
        ("Dangerous Manual Phone Operation While Riding", "Taking hands off handlebars to trigger SOS or mute comms causes crashes.", "Voice-Command Intercom Engine (`voice.js`) providing continuous hands-free speech recognition.")
    ]

    for r_idx, row in enumerate(problems_data, start=1):
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = prob_table.cell(r_idx, c_idx)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True

    # =========================================================================
    # 2. ARCHITECTURE & DIRECTORY TOPOLOGY
    # =========================================================================
    add_h1("2. Software Architecture & Directory Topology")
    
    add_p("The project follows a clean, decoupled, modular component architecture. The HTML file (`index.html`) serves as the application view container; CSS (`styles.css`) defines a high-contrast Deep Space Glassmorphism design system; and 11 focused JavaScript modules encapsulate specialized domain engines.")

    add_h2("2.1 File Directory Structure")
    
    dir_table = doc.add_table(rows=14, cols=4)
    dir_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(dir_table, "CBD5E1")
    
    dir_headers = ["File Path", "Size", "Primary Domain Responsibility", "Exported Symbol / Class"]
    for c_idx, h_text in enumerate(dir_headers):
        cell = dir_table.cell(0, c_idx)
        set_cell_background(cell, "0B132B")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    files_info = [
        ("index.html", "38.5 KB", "Main application UI layout, navigation bar, top header, 5 views, 6 modals & overlays.", "DOM Document"),
        ("css/styles.css", "17.0 KB", "Design system, HSL color tokens, glassmorphism filters, HUD styles, keyframes.", "CSS Stylesheet"),
        ("js/storage.js", "4.6 KB", "LocalStorage persistence manager, default dataset initializers, static CRUD handlers.", "StorageManager"),
        ("js/sound.js", "4.2 KB", "Web Audio API sound synthesizer engine (pings, clicks, chimes, radio beeps, sirens).", "SoundEngine"),
        ("js/voice.js", "4.5 KB", "Web Speech API voice command parser, hands-free intercom state manager.", "VoiceIntercomEngine"),
        ("js/sos.js", "4.8 KB", "Emergency response engine, 5s countdown timer, pitch-sweep siren, multi-channel dispatch.", "SOSEngine"),
        ("js/guardian.js", "4.5 KB", "Timed safety check-in manager, grace period countdown, timeout SOS escalation.", "GuardianEngine"),
        ("js/hazards.js", "8.9 KB", "IndexedDB hazard storage, geotagged map markers, community hazard feed, mesh broadcast.", "HazardEngine"),
        ("js/garage.js", "9.6 KB", "Vehicle maintenance tracker, progress health algorithms, expense logbook, pre-ride check.", "DigitalGarage"),
        ("js/group.js", "6.0 KB", "Ride group host/join manager, access code generator, custom SVG QR code renderer.", "RideGroupEngine"),
        ("js/mesh.js", "9.3 KB", "Off-grid P2P mesh network protocol simulator, Canvas 2D topology graph & particle renderer.", "MeshNetworkEngine"),
        ("js/tracking.js", "24.9 KB", "GPS telemetry tracker, Leaflet map tiles, OSRM turn-by-turn nav, lean angle & crash sensor, GPX export.", "BikeTrackingEngine"),
        ("js/app.js", "21.5 KB", "Main application orchestrator, sub-engine instantiation, navigation router, modals, toast system.", "RidePulseApp")
    ]

    for r_idx, row in enumerate(files_info, start=1):
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = dir_table.cell(r_idx, c_idx)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx in [0, 3]:
                r.bold = True

    add_h2("2.2 System Inter-Dependency & Global Event Bus")
    add_p("All major JavaScript engines attach their instantiated singletons directly to the browser's global `window` object during initialization in `app.js`. This enables seamless cross-module communication without coupling:")
    add_p("• window.app: Instance of RidePulseApp (Central Router & UI Manager)\n• window.digitalGarage: Instance of DigitalGarage (Garage & Maintenance Tracker)\n• window.rideGroup: Instance of RideGroupEngine (Group Hub & SVG QR Code)\n• window.trackingEngine: Instance of BikeTrackingEngine (GPS Telemetry & Maps)\n• window.meshEngine: Instance of MeshNetworkEngine (Canvas 2D P2P Mesh Protocol)\n• window.voiceEngine: Instance of VoiceIntercomEngine (Web Speech Voice Comms)\n• window.sosEngine: Instance of SOSEngine (Emergency Response & Siren Synthesizer)\n• window.guardianEngine: Instance of GuardianEngine (Safety Check-In & Timeout System)\n• window.hazardEngine: Instance of HazardEngine (IndexedDB Geotagged Hazards)\n• window.soundEngine: Instance of SoundEngine (Web Audio Synthesizer)")

    doc.add_page_break()

    # =========================================================================
    # 3. EXHAUSTIVE FILE-BY-FILE & FUNCTION-BY-FUNCTION INVENTORY
    # =========================================================================
    add_h1("3. Exhaustive File-by-File & Function-by-Function Reference")
    add_p("This section provides a detailed technical analysis of every single file in the project, documenting every class, function, method, algorithm, data structure, parameter, return value, and side effect.")

    # 3.1 storage.js
    add_h2("3.1 js/storage.js — LocalStorage Persistence Engine")
    add_p("The `storage.js` module provides centralized persistent storage capabilities utilizing the HTML5 Web Storage API (LocalStorage). It defines fallback defaults for new users and handles serializing and deserializing JSON objects.")
    
    add_h3("Global Constants & Configuration Schema")
    add_p("• STORAGE_KEYS: Map containing keys: RIDER ('ridepulse_rider'), GARAGE ('ridepulse_garage'), ACTIVE_BIKE ('ridepulse_active_bike'), RIDE_LOGS ('ridepulse_ride_logs'), SERVICE_LOGS ('ridepulse_service_logs'), GUARDIAN ('ridepulse_guardian'), SETTINGS ('ridepulse_settings').\n• DEFAULT_RIDER: Object with rider details (Alex Mercer, callsign ApexRider, phone, emergency contacts).\n• DEFAULT_GARAGE: Array of pre-populated bikes (Ducati Panigale V4 S, BMW R 1250 GS Adventure) with mileage intervals.\n• DEFAULT_SERVICE_LOGS: Array of baseline maintenance expense entries.\n• DEFAULT_GUARDIAN: Default safety check-in settings (30m interval, 120s grace period).")

    add_h3("Class: StorageManager (Static Methods)")
    
    st_methods = [
        ("getRider()", "None", "Object", "Reads 'ridepulse_rider' from localStorage. Parses JSON or returns DEFAULT_RIDER if null."),
        ("saveRider(riderData)", "riderData: Object", "void", "Serializes riderData to JSON and writes to localStorage under RIDER key."),
        ("getGarage()", "None", "Array<Object>", "Reads 'ridepulse_garage' from localStorage. Parses JSON array or returns DEFAULT_GARAGE."),
        ("saveGarage(garageData)", "garageData: Array", "void", "Writes array of bike objects to localStorage under GARAGE key."),
        ("getActiveBikeId()", "None", "String", "Returns the active bike ID string (defaults to 'bike_1')."),
        ("setActiveBikeId(bikeId)", "bikeId: String", "void", "Saves bikeId string to localStorage under ACTIVE_BIKE key."),
        ("getActiveBike()", "None", "Object", "Retrieves garage list and active bike ID. Returns matching bike object or garage[0]."),
        ("addBike(newBike)", "newBike: Object", "Object", "Generates unique ID 'bike_' + Date.now(), appends to garage array, saves to storage, sets as active, and returns bike object."),
        ("updateBikeMileage(bikeId, distanceAddedKm)", "bikeId: String, distanceAddedKm: Number", "Object", "Finds bike by ID, increments currentMileage by distanceAddedKm (rounded to 1 decimal), saves garage, returns updated bike."),
        ("getRideLogs()", "None", "Array<Object>", "Reads 'ridepulse_ride_logs' from localStorage. Returns array of saved completed rides or empty array."),
        ("saveRideLog(rideRecord)", "rideRecord: Object", "void", "Unshifts rideRecord to beginning of ride logs array (most recent first) and saves to storage."),
        ("getServiceLogs()", "None", "Array<Object>", "Reads 'ridepulse_service_logs' from localStorage. Returns array or DEFAULT_SERVICE_LOGS."),
        ("saveServiceLog(serviceRecord)", "serviceRecord: Object", "void", "Unshifts serviceRecord to beginning of service logs array and saves to storage."),
        ("getGuardianSettings()", "None", "Object", "Reads 'ridepulse_guardian' from localStorage or returns DEFAULT_GUARDIAN."),
        ("saveGuardianSettings(settings)", "settings: Object", "void", "Saves guardian settings object to localStorage under GUARDIAN key.")
    ]

    for m_name, m_params, m_ret, m_desc in st_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• StorageManager.{m_name}\n")

    # 3.2 sound.js
    add_h2("3.2 js/sound.js — Web Audio Synthesizer Engine")
    add_p("The `sound.js` module encapsulates pure Web Audio API sound synthesis, eliminating the need for external MP3/WAV audio assets. It creates custom audio nodes on demand to produce tactical UI clicks, inter-node mesh pings, radio chimes, and dual-tone emergency sirens.")

    add_h3("Class: SoundEngine Methods")
    
    snd_methods = [
        ("initContext()", "None", "void", "Instantiates window.AudioContext or webkitAudioContext if null. Resumes context if state is 'suspended'."),
        ("playClick()", "None", "void", "Generates a 50ms sine wave frequency sweep (800Hz to 400Hz) with exponential gain decay for tactile UI button clicks."),
        ("playChime()", "None", "void", "Plays a 4-note ascending major arpeggio [C5 (523.25Hz), E5 (659.25Hz), G5 (783.99Hz), C6 (1046.50Hz)] using staggered sine oscillators for milestone notifications."),
        ("playMeshPing()", "None", "void", "Generates a high-pitched 150ms sine sweep (1200Hz to 1800Hz) representing a successful P2P mesh packet hop."),
        ("playRadioBeep()", "None", "void", "Generates a 80ms 950Hz square wave tone mimicking tactical walkie-talkie radio mic click feedback."),
        ("startSirenAlert()", "None", "void", "Stops existing siren if active. Creates a sawtooth oscillator and gains node. Starts a 280ms interval alternating pitch between 650Hz and 950Hz for emergency alerts."),
        ("stopSirenAlert()", "None", "void", "Clears siren interval timer, stops sawtooth oscillator, disconnects gain node, and resets siren instance.")
    ]

    for m_name, m_params, m_ret, m_desc in snd_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• SoundEngine.{m_name}\n")

    # 3.3 voice.js
    add_h2("3.3 js/voice.js — Voice-Command Intercom Engine")
    add_p("The `voice.js` module implements continuous hands-free voice recognition using the browser's Web Speech API (`SpeechRecognition`). It allows riders to trigger emergency SOS alerts, toggle intercom mic mute state, or query ride telemetry without taking their hands off the handlebars.")

    add_h3("Class: VoiceIntercomEngine Methods")
    
    voice_methods = [
        ("init(onCommandCallback)", "onCommandCallback: Function", "void", "Registers command handler callback and calls setupSpeechRecognition()."),
        ("setupSpeechRecognition()", "None", "void", "Instantiates SpeechRecognition/webkitSpeechRecognition. Sets continuous=true, interimResults=false, lang='en-US'. Attaches onresult, onerror, and onend event listeners for continuous listening."),
        ("startListening()", "None", "void", "Sets isListening=true, calls recognition.start(), initializes Web Audio context, and updates top bar VOX pill UI."),
        ("stopListening()", "None", "void", "Sets isListening=false, calls recognition.stop(), and updates VOX pill UI to standby mode."),
        ("toggleMicMute()", "None", "boolean", "Toggles isMicMuted boolean state. Plays high/low Web Audio feedback tone (440Hz/880Hz), updates VOX pill UI, and returns new mute state."),
        ("parseVoiceCommand(transcript)", "transcript: String", "void", "Parses spoken phrase. Matches 'sos'/'emergency' -> SOS_TRIGGER; 'mute' -> MUTE_MIC; 'unmute' -> UNMUTE_MIC; 'status'/'speed' -> CHECK_STATUS. Invokes onCommandCallback."),
        ("initAudioContext()", "None", "void", "Creates fallback AudioContext for audio feedback tones."),
        ("playTone(frequency, duration)", "frequency: Number, duration: Number", "void", "Synthesizes a short sine tone for voice intercom status changes."),
        ("updateVoxPillUI()", "None", "void", "Updates DOM element '#vox-status-pill' with active, muted, or standby CSS classes and icons.")
    ]

    for m_name, m_params, m_ret, m_desc in voice_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• VoiceIntercomEngine.{m_name}\n")

    # 3.4 sos.js
    add_h2("3.4 js/sos.js — SOS Emergency Response Engine")
    add_p("The `sos.js` module handles critical emergency response workflows. When triggered manually or automatically (via Guardian timeout or crash sensor), it initiates a 5-second safety countdown, triggers a red strobe overlay, plays a pitch-sweep emergency siren, builds an emergency payload, and dispatches it via Cellular API and Off-Grid Mesh Relays.")

    add_h3("Class: SOSEngine Methods")
    
    sos_methods = [
        ("triggerSOS(reason, sosType)", "reason: String, sosType: String", "void", "Activates emergency mode (countdownSeconds=5, isSOSActive=true). Renders '#sos-overlay' with strobe animation, starts siren sound, and begins 1-second countdown timer interval. Calls dispatchSOSPayload() upon reaching 0."),
        ("cancelSOS()", "None", "void", "Aborts emergency countdown, clears timer, stops siren sound, hides overlay UI, and broadcasts notification."),
        ("dispatchSOSPayload(reason, sosType)", "reason: String, sosType: String", "void", "Gathers rider info, active bike details, coordinates, and emergency contacts. Formats JSON payload. Relays packet via `window.meshEngine.sendPacket('SOS', payload)`. Updates status UI with dispatch confirmation."),
        ("updateCountdownUI()", "None", "void", "Updates '#sos-countdown-val' DOM element with remaining countdown seconds."),
        ("startSirenSound()", "None", "void", "Creates sawtooth AudioContext oscillator with linear frequency ramping between 600Hz and 1200Hz every 0.25 seconds."),
        ("stopSirenSound()", "None", "void", "Stops sawtooth oscillator and closes siren AudioContext.")
    ]

    for m_name, m_params, m_ret, m_desc in sos_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• SOSEngine.{m_name}\n")

    # 3.5 guardian.js
    add_h2("3.5 js/guardian.js — Guardian Check-In Engine")
    add_p("The `guardian.js` module provides automated rider safety check-ins. During a ride, it triggers periodic prompts at configurable intervals (e.g. 30 minutes). If the rider fails to acknowledge 'I'M OK!' within the grace period (e.g. 2 minutes), it automatically escalates and triggers emergency SOS dispatch.")

    add_h3("Class: GuardianEngine Methods")
    
    gdn_methods = [
        ("init()", "None", "void", "Loads saved settings from StorageManager."),
        ("loadSettings()", "None", "void", "Fetches guardian settings (enabled, intervalMins, graceSec, contactName, contactPhone) from StorageManager."),
        ("saveSettings(newSettings)", "newSettings: Object", "void", "Updates instance properties and persists settings to StorageManager."),
        ("startRideGuardian()", "None", "void", "Clears existing timers. If enabled, starts setInterval check-in timer for (intervalMins * 60 * 1000) ms."),
        ("stopRideGuardian()", "None", "void", "Clears check-in and grace period timers and hides prompt UI."),
        ("triggerCheckInPrompt()", "None", "void", "Plays alert chime, displays floating '#guardian-prompt-overlay', and starts 1-second interval countdown for remaining grace seconds."),
        ("acknowledgeCheckIn()", "None", "void", "Rider clicks 'I'M OK!'. Stops grace countdown timer, hides prompt UI, plays chime, and logs safety confirmation."),
        ("handleTimeoutSOS()", "None", "void", "Executed when grace countdown reaches 0 without response. Hides prompt UI, starts siren alert, and triggers SOSEngine with 'guardian_timeout' payload flag."),
        ("hidePromptUI()", "None", "void", "Removes active CSS class from '#guardian-prompt-overlay'."),
        ("updateGraceCountdownUI()", "None", "void", "Formats remaining grace seconds into MM:SS and writes to '#guardian-countdown-val'.")
    ]

    for m_name, m_params, m_ret, m_desc in gdn_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• GuardianEngine.{m_name}\n")

    doc.add_page_break()

    # 3.6 hazards.js
    add_h2("3.6 js/hazards.js — Community Hazard Reporting Engine")
    add_p("The `hazards.js` module manages geotagged road hazard reports (potholes, oil spills, crash scenes, signal blind spots, wildlife). It stores data natively in browser IndexedDB (`RidePulse_DB`) with LocalStorage fallback, renders Leaflet map markers, updates the hazard feed UI, and broadcasts reports over the P2P Mesh Network.")

    add_h3("Data Configuration: HAZARD_CATEGORIES")
    add_p("• POTHOLE: Pothole / Road Damage (Icon: fa-road-spikes, Color: #ff9100)\n• DEBRIS: Debris / Oil Spill / Gravel (Icon: fa-triangle-exclamation, Color: #ffab00)\n• ACCIDENT: Traffic Crash / Stalled Bike (Icon: fa-car-burst, Color: #ff1744)\n• SIGNAL_DOWN: Blind Corner / Signal Down (Icon: fa-wifi, Color: #d500f9)\n• ANIMAL: Wildlife / Animal on Road (Icon: fa-paw, Color: #00e676)")

    add_h3("Class: HazardEngine Methods")
    
    hzd_methods = [
        ("init()", "Async", "void", "Initializes IndexedDB store. Loads saved hazards. Seeds demo sample hazards if empty. Renders feed UI and map markers."),
        ("initIndexedDB()", "Async", "Promise<IDBDatabase>", "Opens IndexedDB 'RidePulse_DB' v1, creates object store 'hazards' with keyPath 'id'."),
        ("loadHazards()", "Async", "void", "Queries IndexedDB object store (or fallback LocalStorage) to fetch all hazards. Sorts recency descending."),
        ("saveHazardToStorage(hazard)", "Async", "void", "Puts hazard record into IndexedDB object store or saves to LocalStorage array."),
        ("reportHazard(categoryKey, note, coords)", "Async", "Object", "Constructs hazard object (id, category, icon, color, note, reporter, lat, lng, timestamp, confirmations=1). Saves to storage, broadcasts via MeshEngine.sendPacket('HAZARD'), adds Leaflet marker, updates feed UI, and returns hazard object."),
        ("confirmHazard(hazardId)", "Async", "void", "Increments confirmation count for hazard ('Still There'), updates storage, map popup, and feed UI."),
        ("seedDemoHazards()", "Async", "void", "Populates initial demo hazards (Deep pothole, Spilled diesel fuel, Deer crossing) if local database is empty."),
        ("plotAllHazardsOnMap()", "None", "void", "Iterates hazardsList and calls window.trackingEngine.addHazardMarker(h)."),
        ("renderHazardFeedUI()", "None", "void", "Generates HTML feed cards with category icons, recency timestamps, reporter callsign, and confirmation button.")
    ]

    for m_name, m_params, m_ret, m_desc in hzd_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• HazardEngine.{m_name}\n")

    # 3.7 garage.js
    add_h2("3.7 js/garage.js — Digital Garage & Maintenance Engine")
    add_p("The `garage.js` module provides complete fleet management for multi-bike owners. It calculates real-time service item health percentages, predicts overdue maintenance, enforces pre-ride safety gates, and maintains a maintenance expense logbook tracking cost-per-kilometer.")

    add_h3("Class: DigitalGarage Methods")
    
    gar_methods = [
        ("init()", "None", "void", "Calls refresh() to initialize garage UI."),
        ("refresh()", "None", "void", "Fetches garage list and active bike from StorageManager. Calls renderGarage(), renderActiveBikeSummary(), and renderServiceLogbook()."),
        ("calculateItemHealth(currentMileage, lastDoneMileage, intervalKm)", "Static", "Object", "Algorithm: kmDriven = current - lastDone; kmRemaining = interval - kmDriven; percentRemaining = clamp((kmRemaining / interval) * 100). Status: 'OVERDUE' if remaining <= 0; 'DUE SOON' if percent <= 20; else 'GOOD'."),
        ("getBikeMaintenanceOverview(bike)", "None", "Object", "Evaluates health for 4 core items: Engine Oil, Brake Pads, Chain Lubrication, and Tire Tread. Returns item array, overdueCount, and warningCount."),
        ("renderGarage()", "None", "void", "Renders vehicle cards in '#garage-bikes-container' with progress bars, status badges, set active buttons, and reset service controls."),
        ("renderActiveBikeSummary()", "None", "void", "Updates top header badge '#active-bike-summary-badge' with active bike name, total mileage, and overdue count."),
        ("renderServiceLogbook()", "None", "void", "Renders expense history table in '#service-logbook-table-body'. Calculates total spent ($) and cost per kilometer ($/km)."),
        ("addServiceExpense(type, cost, mileage, notes)", "None", "void", "Constructs service log record, saves via StorageManager, plays chime, and refreshes UI."),
        ("selectActiveBike(bikeId)", "bikeId: String", "void", "Sets new active bike ID in StorageManager and refreshes garage UI."),
        ("markServicedModal(bikeId)", "bikeId: String", "void", "Resets bike service mileage checkpoints (oil, brakes, chain, tires) to current mileage. Saves to storage and refreshes UI."),
        ("addNewBike(make, model, year, currentMileage)", "None", "Object", "Adds new bike record to StorageManager with default interval thresholds and sets it as active.")
    ]

    for m_name, m_params, m_ret, m_desc in gar_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• DigitalGarage.{m_name}\n")

    # 3.8 group.js
    add_h2("3.8 js/group.js — Ride Group & SVG QR Code Invite Engine")
    add_p("The `group.js` module facilitates crew ride formation. It generates random 6-character access codes, renders dynamic inline SVG QR codes without external graphic libraries, tracks member ready status, and displays active group rosters.")

    add_h3("Class: RideGroupEngine Methods")
    
    grp_methods = [
        ("init()", "None", "void", "Loads baseline crew members (Alex, Dave R., Maya S.) and access code 'PULSE9'."),
        ("createGroup(groupName)", "groupName: String", "Object", "Generates access code via generateAccessCode(), sets currentGroup object, and calls renderGroupUI()."),
        ("joinGroup(code)", "code: String", "boolean", "Validates input code string, updates active access code, renders UI, and returns success boolean."),
        ("generateAccessCode()", "None", "String", "Generates 6-character alphanumeric code using characters 'ABCDEFGHJKLMNPQRSTUVWXYZ23456789'."),
        ("renderGroupUI()", "None", "void", "Writes code to DOM '#group-code-display', calls drawSVGQRCode(), and renders roster list in '#group-roster-container'."),
        ("drawSVGQRCode(container, text)", "container: HTMLElement, text: String", "void", "Generates dynamic inline SVG markup containing corner alignment markers, data matrix rects, and cyan glow accents.")
    ]

    for m_name, m_params, m_ret, m_desc in grp_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• RideGroupEngine.{m_name}\n")

    # 3.9 mesh.js
    add_h2("3.9 js/mesh.js — Off-Grid Mesh Network Engine")
    add_p("The `mesh.js` module simulates P2P hop-by-hop packet relay across crew devices when cellular signal is lost. It renders an animated 2D Canvas graph depicting multi-hop packet routing (Self -> Dave -> Maya -> Cell Gateway -> Cloud).")

    add_h3("Class: MeshNetworkEngine Methods")
    
    mesh_methods = [
        ("init(canvasId)", "canvasId: String", "void", "Attaches to HTML5 Canvas element, initializes 2D context, attaches window resize listener, and starts animation loop."),
        ("setCellularStatus(isOnline)", "isOnline: boolean", "void", "Updates isCellularOnline boolean state and top bar signal status pill UI."),
        ("toggleCellular()", "None", "boolean", "Toggles cellular connection and returns new online state."),
        ("sendPacket(packetType, payload)", "packetType: String, payload: Object", "Object", "If cellular online, transmits direct to cloud (1 hop). If cellular offline, initiates multi-hop P2P animation (Self -> Dave -> Maya -> Gateway) and logs packet delivery."),
        ("triggerPacketAnimation(fromIndex, toIndex, onComplete)", "fromIndex: Number, toIndex: Number, onComplete: Function", "void", "Pushes animated packet particle object into activePackets array with progress interpolation."),
        ("getNodeCanvasPos(index)", "index: Number", "Object", "Returns relative canvas (x, y) coordinates for topology node."),
        ("logPacket(packet, note)", "packet: Object, note: String", "void", "Unshifts packet event into packetLog array and calls renderPacketLogUI()."),
        ("renderPacketLogUI()", "None", "void", "Renders real-time telemetry log lines in '#mesh-packet-logs'."),
        ("resizeCanvas()", "None", "void", "Resizes canvas element to match container bounding rectangle."),
        ("startCanvasAnimation()", "None", "void", "Starts requestAnimationFrame loop executing drawTopologyGraph()."),
        ("drawTopologyGraph()", "None", "void", "Draws background grid, dashed connection links, node circles, pulse aura rings, labels, and animated packet dots.")
    ]

    for m_name, m_params, m_ret, m_desc in mesh_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• MeshNetworkEngine.{m_name}\n")

    doc.add_page_break()

    # 3.10 tracking.js
    add_h2("3.10 js/tracking.js — Bike Experience Tracking & Navigation Engine")
    add_p("The `tracking.js` module is the largest core engine. It manages real-time GPS location tracking, Leaflet interactive map rendering, multiple map tile layers (Google Maps Streets, Satellite, Google Hybrid, Dark Mode, OSM), Google Live Traffic overlay, OSRM turn-by-turn navigation routing, lean angle calculation via device gyroscope, G-force estimation, simulated crash impact sensors, and GPX track export.")

    add_h3("Class: BikeTrackingEngine Methods")
    
    trk_methods = [
        ("constructor()", "None", "void", "Initializes telemetry counters (distance, speed, top speed, moving time, lean angles, G-force), map tile URLs, and demo simulated twisty route coordinates."),
        ("setupOrientationListener()", "None", "void", "Attaches listener to window.deviceorientation event. Reads gamma property to calculate bike lean angle in real-time."),
        ("initMap(mapElementId)", "mapElementId: String", "void", "Instantiates Leaflet map centered on current coordinates. Adds Google Maps tile layer, polyline layer, custom rider marker, and group peer markers."),
        ("locateUserGPS()", "None", "void", "Calls navigator.geolocation.getCurrentPosition() with high accuracy. Updates rider marker position, pans map, and displays accuracy radius."),
        ("searchLocation(query)", "Async", "void", "Queries OpenStreetMap Nominatim Geocoding API for destination query string. Pans Leaflet map to result coordinates."),
        ("invalidateMapSize()", "None", "void", "Triggers map.invalidateSize() with timeout delays to re-calculate container bounds after tab/view switching."),
        ("toggleTraffic()", "None", "void", "Toggles Google Maps Live Traffic tile layer overlay ('mt1.google.com/vt/lyrs=m,traffic'). Updates traffic toggle button UI."),
        ("calculateRouteToDestination(destinationQuery)", "Async", "void", "Geocodes destination, fetches OSRM driving route geometry, draws Google Blue navigation route line, places destination marker, and displays turn-by-turn banner UI with ETA and distance."),
        ("cancelNavigation()", "None", "void", "Removes navigation polyline and destination marker from map and hides turn-by-turn banner UI."),
        ("switchMapLayer(styleName)", "styleName: String", "void", "Switches map tile layer between 'google', 'googleSat', 'googleHybrid', 'dark', and 'osm'."),
        ("addPeerMarker(id, name, coords, color)", "id: String, name: String, coords: Array, color: String", "void", "Creates custom divIcon marker for group crew members and adds to Leaflet map."),
        ("addHazardMarker(hazard)", "hazard: Object", "void", "Creates custom divIcon marker with hazard category icon and color. Binds interactive popup with 'Still There' confirmation button."),
        ("startTracking()", "None", "void", "Resets session telemetry, starts 1-second moving timer, begins navigator.geolocation.watchPosition(), and initiates desktop twisty route simulator."),
        ("pauseTracking()", "None", "boolean", "Toggles tracking pause state and returns isPaused boolean."),
        ("stopTracking()", "None", "Object", "Stops timers and GPS watch, constructs rideRecord object, saves to StorageManager, updates bike odometer, and returns rideRecord."),
        ("handleGPSUpdate(lat, lng, speedMps, ele)", "lat: Number, lng: Number, speedMps: Number, ele: Number", "void", "Calculates speed (km/h), top speed, G-force, and Haversine incremental distance. Appends point to routePolyline, updates rider marker position, and broadcasts mesh location updates."),
        ("updateLeanAngle(angle)", "angle: Number", "void", "Clamps lean angle between -50° and +50°. Updates max lean left/right records and calls updateLeanAngleUI()."),
        ("updateLeanAngleUI()", "None", "void", "Rotates bike tilt DOM graphic ('#hud-bike-tilt-graphic') and updates angle and G-force text."),
        ("triggerSimulatedCrash()", "None", "void", "Simulates high-G crash impact (3.8 G spike). Plays siren alert and triggers emergency SOS dispatch."),
        ("exportGPX(rideRecord)", "rideRecord: Object", "void", "Generates valid GPX 1.1 XML document string containing trackpoints with latitude, longitude, elevation, and timestamp. Triggers browser file download."),
        ("haversineDistance(lat1, lon1, lat2, lon2)", "Static", "Number", "Calculates great-circle distance in kilometers between two GPS coordinates using the Haversine formula."),
        ("updateHUDTimerUI()", "None", "void", "Formats moving seconds to MM:SS and updates '#hud-timer-val'."),
        ("updateHUDValuesUI()", "None", "void", "Updates speedometer, distance, and top speed text elements in telemetry HUD.")
    ]

    for m_name, m_params, m_ret, m_desc in trk_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• BikeTrackingEngine.{m_name}\n")

    # 3.11 app.js
    add_h2("3.11 js/app.js — Main Application Orchestrator")
    add_p("The `app.js` module serves as the central control nexus. It instantiates all singletons, wires up DOM navigation tab switching, registers button event listeners, manages modal open/close transitions, handles incoming voice commands, renders post-ride summary modals, and displays toast notification alerts.")

    add_h3("Class: RidePulseApp Methods")
    
    app_methods = [
        ("init()", "None", "void", "Instantiates all 9 core engine classes, executes initialization sequences, sets up navigation router and event listeners, renders ride logs, initializes Leaflet map, and displays welcome toast."),
        ("setupNavigation()", "None", "void", "Attaches click event listeners to navigation bar items carrying 'data-view' attributes."),
        ("switchView(viewId)", "viewId: String", "void", "Removes active class from all view sections and nav items. Activates target view section ('#garage-view', '#group-view', '#tracking-view', '#mesh-view', '#logs-view') and triggers view-specific refresh logic."),
        ("setupEventListeners()", "None", "void", "Registers click/change listeners for modal buttons, garage pre-ride check, group host/join buttons, cellular toggle switch, mic mute, SOS trigger, crash sensor test, hazard category picker, service log saver, and GPX export."),
        ("handleVoiceCommand(cmdType, transcript)", "cmdType: String, transcript: String", "void", "Processes voice commands forwarded by VoiceIntercomEngine. Triggers SOS, mutes mic, or displays status toast."),
        ("renderPostRideModal(rideRecord)", "rideRecord: Object", "void", "Constructs HTML summary for completed ride (distance, duration, top speed, max lean, garage sync confirmation) and updates modal DOM body."),
        ("renderRideLogs()", "None", "void", "Reads ride logs array from StorageManager and generates HTML log cards with GPX export buttons in '#ride-logs-list'."),
        ("openModal(modalId)", "modalId: String", "void", "Adds 'active' class to modal backdrop element to display popup window."),
        ("closeModal(modalId)", "modalId: String", "void", "Removes 'active' class from modal backdrop element to close window."),
        ("showNotification(msg, type)", "msg: String, type: String", "void", "Creates animated toast notification div with dynamic border colors (info: cyan, success: green, warning: amber, danger: crimson), appends to '#toast-container', and auto-removes after 4 seconds.")
    ]

    for m_name, m_params, m_ret, m_desc in app_methods:
        add_p(f"Signature: {m_name}\nParameters: {m_params} | Returns: {m_ret}\nDescription: {m_desc}", bold_prefix=f"• RidePulseApp.{m_name}\n")

    doc.add_page_break()

    # =========================================================================
    # 4. KEY FEATURES & FUNCTIONAL WORKFLOWS
    # =========================================================================
    add_h1("4. Key Features & Functional Workflows")

    add_h2("4.1 Feature 1: Digital Garage & Predictive Maintenance")
    add_p("The Digital Garage allows riders to maintain multiple motorbikes (e.g., Ducati Panigale V4 S, BMW R 1250 GS). Each bike maintains mileage intervals for 4 critical service items: Engine Oil, Brake Pads/Fluid, Chain Lubrication/Tension, and Tire Tread/Pressure. The `calculateItemHealth()` algorithm calculates the percentage of remaining useful life and assigns health statuses ('GOOD', 'DUE SOON', 'OVERDUE'). When starting a ride, the Pre-Ride Health Check acts as a safety gate—if any service item is overdue, a modal warning prompts the rider before allowing entry to the Ride Hub.")

    add_h2("4.2 Feature 2: Group Ride Hub & SVG QR Invites")
    add_p("To eliminate tedious manual phone number sharing, RidePulse enables trip hosts to generate an instant crew session with a 6-character access code (e.g. 'PULSE9'). The app dynamically generates an inline SVG QR code without external dependencies. Peer riders can scan or enter the code to join the live roster, which displays real-time connection status (5G Cellular vs. Mesh Hop Node) and ready status.")

    add_h2("4.3 Feature 3: Live Telemetry HUD, Lean Angle & Crash Sensor")
    add_p("During active GPS tracking, the Live HUD displays high-contrast telemetry: Speedometer (km/h), Distance (km), Duration (MM:SS), Top Speed, Elevation, and G-Force. The app integrates with the device gyroscope (`deviceorientation` event) to measure bank/lean angles in real-time, displaying a dynamic tilting motorcycle graphic. In the event of a high-G impact (e.g., 3.8 G spike), the Crash Impact Sensor automatically triggers an emergency SOS dispatch.")

    add_h2("4.4 Feature 4: Off-Grid P2P Mesh Network & Hop Relay")
    add_p("When cellular connectivity is lost in mountainous or remote terrain, the rider can toggle off cellular signal. RidePulse seamlessly transitions to Off-Grid Mesh Mode. Utilizing an HTML5 Canvas 2D topology graph renderer, the app visualizes hop-by-hop packet relay across crew devices (Rider A -> Rider B -> Rider C -> Cell Tower Gateway). Packets (location updates, hazard alerts, SOS dispatches) hop across peer nodes until reaching a node with active cellular uplink.")

    add_h2("4.5 Feature 5: Voice Intercom & Synthesized Web Audio")
    add_p("Riders can activate continuous hands-free voice control using the Web Speech API. Spoken commands such as 'SOS Emergency', 'Mute Intercom', or 'Check Status' are parsed in real-time. Sound effects (clicks, chimes, radio beeps, pitch-sweep sirens) are synthesized dynamically via the Web Audio API without relying on external sound assets.")

    add_h2("4.6 Feature 6: Automated Guardian Safety Check-In")
    add_p("For solo riders, the Guardian system prompts the user at regular intervals (e.g. 30 minutes) with a floating visual prompt and chime. If the rider acknowledges by clicking 'I'M OK!', the countdown resets. If the rider is unresponsive after the grace period (e.g. 2 minutes), Guardian automatically escalates to full SOS Emergency mode.")

    add_h2("4.7 Feature 7: Community Hazard Reporting & IndexedDB")
    add_p("Riders can report geotagged road hazards (Potholes, Oil Spills, Accidents, Signal Blind Spots, Wildlife). Reports are stored locally in IndexedDB (`RidePulse_DB`), plotted on the Leaflet map with custom icons, broadcast to nearby crew members via the mesh network, and rendered in an interactive hazard feed supporting crowd-sourced 'Still There' confirmations.")

    add_h2("4.8 Feature 8: Emergency SOS Response System")
    add_p("Triggering SOS activates a red strobe backdrop overlay, pitch-sweep siren audio, and a 5-second cancelable countdown timer. Upon completion, an emergency payload containing GPS coordinates, bike specs, rider details, and emergency contacts is dispatched via both cellular API and P2P mesh relay.")

    add_h2("4.9 Feature 9: Multi-Layer Google Navigation & Live Traffic")
    add_p("The map card integrates Leaflet with multiple map tile providers: Google Maps Streets, Satellite, Google Hybrid, Dark Mode, and OpenStreetMap. It features a toggleable Google Live Traffic overlay and OSRM turn-by-turn navigation routing with real-time ETA, distance, and turn instruction banners.")

    add_h2("4.10 Feature 10: Ride Analytics & GPX Track Export")
    add_p("Upon completing a ride, total distance is automatically added to the active bike's odometer in the Digital Garage. Ride logs are saved in history, and riders can export their full GPS route track as a standard GPX 1.1 XML file for analysis in third-party mapping software.")

    doc.add_page_break()

    # =========================================================================
    # 5. DATA MODELS & STORAGE SCHEMAS
    # =========================================================================
    add_h1("5. Data Models & Storage Schemas")
    
    add_h2("5.1 LocalStorage Keys & Data Structure Schemas")
    
    add_p("1. Garage Vehicle Object Schema (`ridepulse_garage`):", bold_prefix="• ")
    doc.add_paragraph("{\n  id: 'bike_1',\n  make: 'Ducati',\n  model: 'Panigale V4 S',\n  year: 2023,\n  currentMileage: 4850,\n  lastOilChangeMileage: 4200, oilIntervalKm: 3000,\n  lastBrakeCheckMileage: 3500, brakeIntervalKm: 2500,\n  lastChainLubeMileage: 4600, chainIntervalKm: 500,\n  lastTireCheckMileage: 3000, tireIntervalKm: 4000\n}")

    add_p("2. Ride Record Schema (`ridepulse_ride_logs`):", bold_prefix="• ")
    doc.add_paragraph("{\n  id: 'ride_1700000000000',\n  date: 'Sat, Aug 1, 2026',\n  time: '11:15:00 AM',\n  durationSec: 1450,\n  distanceKm: 42.5,\n  avgSpeedKmh: 75,\n  topSpeedKmh: 115,\n  maxLeanLeft: 35,\n  maxLeanRight: 41,\n  maxGForce: 1.35,\n  bikeId: 'bike_1',\n  bikeName: 'Ducati Panigale V4 S',\n  routePoints: [{ lat: 37.7749, lng: -122.4194, speed: 42, ele: 45 }, ...]\n}")

    add_h2("5.2 IndexedDB Object Store Schema (`RidePulse_DB`)")
    add_p("Database: `RidePulse_DB` | Store Name: `hazards` | KeyPath: `id`")
    doc.add_paragraph("{\n  id: 'hzd_1700000000000',\n  categoryKey: 'POTHOLE',\n  categoryName: 'Pothole / Road Damage',\n  icon: 'fa-road-spikes',\n  color: '#ff9100',\n  note: 'Deep pothole on hairpin turn apex',\n  reporter: 'Dave R. (KTM 890)',\n  lat: 37.7772,\n  lng: -122.4235,\n  timestamp: 1700000000000,\n  dateStr: '15 mins ago',\n  confirmations: 3\n}")

    # =========================================================================
    # 6. CONCLUSION & ROADMAP
    # =========================================================================
    add_h1("6. System Verification & Future Roadmap")
    
    add_p("RidePulse successfully fulfills all requirements for a state-of-the-art bike experience tracking app. The entire solution operates reliably in modern web browsers without backend server setup.")

    add_h2("6.1 Future Enhancement Roadmap")
    add_p("1. Bluetooth Low Energy (BLE) Mesh Integration: Connect physical Web Bluetooth microcontrollers (ESP32/nRF52) for true hardware RF mesh relays.\n2. OBD-II Bluetooth ECU Integration: Read real-time engine RPM, throttle position, lean angle, and coolant temperature directly from motorcycle CAN-bus systems.\n3. Vector Offline Map Tile Caching: Pre-cache Leaflet map vector tiles via Progressive Web App (PWA) Service Workers for 100% offline map rendering.")

    # Save document
    output_filename = "RidePulse Technical Project Report.docx"
    doc.save(output_filename)
    print(f"Report document successfully generated and saved to: {output_filename}")

if __name__ == '__main__':
    generate_report()
